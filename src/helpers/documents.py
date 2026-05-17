from io import BytesIO
from typing import Protocol, cast
from zipfile import BadZipFile

import pandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from odf.opendocument import load as odf_load
from odf.text import H, P

import discord
import httpx

from .ui import (
    SUPPORTED_WORD_ATTACHMENT_EXTENSIONS,
    SUPPORTED_WORD_CONTENT_TYPES,
)

DOCUMENT_FORMAT_CHOICES = [
    discord.app_commands.Choice(name="Markdown (arquivo .md)", value="md"),
    discord.app_commands.Choice(name="DOCX (Microsoft Word)", value="docx"),
    discord.app_commands.Choice(name="ODT (LibreOffice)", value="odt"),
    discord.app_commands.Choice(name="PDF", value="pdf"),
]


class DocumentProcessor(Protocol):
    extension: str

    def extract_text(self, document_bytes: bytes) -> str: ...
    def generate(self, content: str, title: str) -> bytes: ...


class DocxProcessor:
    extension: str = ".docx"

    def extract_text(self, document_bytes: bytes) -> str:
        try:
            doc = Document(BytesIO(document_bytes))
        except (BadZipFile, KeyError) as exc:
            raise ValueError("invalid_docx") from exc

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def generate(self, content: str, title: str) -> bytes:
        docx_bytes = _run_pandoc(content, "docx")
        docx_bytes = self._apply_abnt(docx_bytes)
        return self._insert_centered_title(docx_bytes, title)

    def _insert_centered_title(self, docx_bytes: bytes, title: str) -> bytes:
        doc = Document(BytesIO(docx_bytes))

        if doc.paragraphs:
            title_paragraph = doc.paragraphs[0].insert_paragraph_before(title)
        else:
            title_paragraph = doc.add_paragraph(title)

        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.style = "Normal"

        for run in title_paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.bold = True

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _apply_abnt(self, docx_bytes: bytes) -> bytes:
        doc = Document(BytesIO(docx_bytes))

        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

        try:
            normal = doc.styles["Normal"]
        except KeyError:
            pass
        else:
            normal.font.name = "Times New Roman"  # pyright: ignore[reportAttributeAccessIssue]
            normal.font.size = Pt(12)  # pyright: ignore[reportAttributeAccessIssue]
            normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # pyright: ignore[reportAttributeAccessIssue]
            normal.paragraph_format.line_spacing = 1.5  # pyright: ignore[reportAttributeAccessIssue]

        for name in (
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "heading 1",
            "heading 2",
            "heading 3",
        ):
            try:
                heading = doc.styles[name]
            except KeyError:
                continue
            heading.font.name = "Times New Roman"  # pyright: ignore[reportAttributeAccessIssue]
            heading.font.size = Pt(12)  # pyright: ignore[reportAttributeAccessIssue]
            heading.font.bold = True  # pyright: ignore[reportAttributeAccessIssue]
            heading.paragraph_format.line_spacing = 1.5  # pyright: ignore[reportAttributeAccessIssue]
            heading.paragraph_format.space_before = Pt(12)  # pyright: ignore[reportAttributeAccessIssue]
            heading.paragraph_format.space_after = Pt(6)  # pyright: ignore[reportAttributeAccessIssue]

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


class OdtProcessor:
    extension: str = ".odt"

    def extract_text(self, document_bytes: bytes) -> str:
        try:
            doc = odf_load(BytesIO(document_bytes))
        except (BadZipFile, KeyError) as exc:
            raise ValueError("invalid_odt") from exc

        paragraphs = []
        for elem in doc.getElementsByType(P):
            text = str(elem).strip()
            if text:
                paragraphs.append(text)
        for elem in doc.getElementsByType(H):
            text = str(elem).strip()
            if text:
                paragraphs.append(text)

        return "\n\n".join(paragraphs)

    def generate(self, content: str, title: str) -> bytes:
        md = f"# {title}\n\n{content}"
        return _run_pandoc(md, "odt")


class MdProcessor:
    extension: str = ".md"

    def extract_text(self, document_bytes: bytes) -> str:
        return document_bytes.decode("utf-8")

    def generate(self, content: str, title: str) -> bytes:
        md = f"# {title}\n\n{content}"
        return md.encode("utf-8")


class PdfProcessor:
    extension: str = ".pdf"

    def extract_text(self, document_bytes: bytes) -> str:
        raise ValueError("pdf_extraction_not_supported")

    def generate(self, content: str, title: str) -> bytes:
        md = f"\\begin{{center}}\n\\textbf{{{title}}}\n\\end{{center}}\n\n{content}"
        return _run_pandoc(md, "pdf")


_PROCESSORS: dict[str, type[DocumentProcessor]] = {
    "docx": DocxProcessor,
    "odt": OdtProcessor,
    "md": MdProcessor,
    "pdf": PdfProcessor,
}


def get_processor(format_or_ext: str) -> DocumentProcessor:
    key = format_or_ext.lower().lstrip(".")
    processor_cls = _PROCESSORS.get(key)
    if processor_cls is None:
        raise ValueError(f"unsupported_document_format: {format_or_ext}")
    return processor_cls()


def attachment_is_supported_word_document(attachment: discord.Attachment) -> bool:
    content_type = attachment.content_type or ""
    filename = attachment.filename.lower()
    return content_type in SUPPORTED_WORD_CONTENT_TYPES or filename.endswith(
        SUPPORTED_WORD_ATTACHMENT_EXTENSIONS
    )


async def read_word_attachment(
    attachment: discord.Attachment, max_chars: int, http_client: httpx.AsyncClient
) -> tuple[str, bool]:
    if not attachment_is_supported_word_document(attachment):
        raise ValueError("unsupported")

    response = await http_client.get(attachment.url)
    response.raise_for_status()

    ext = attachment.filename.lower().rsplit(".", 1)[-1]
    processor = get_processor(ext)
    text = processor.extract_text(response.content)

    return text[:max_chars], len(text) > max_chars


def _run_pandoc(markdown_text: str, output_format: str) -> bytes:
    try:
        doc = pandoc.read(source=markdown_text, format="markdown")
    except RuntimeError as exc:
        raise RuntimeError(
            "pandoc is required to generate documents. Install it from https://pandoc.org/installing.html"
        ) from exc
    try:
        if output_format == "pdf":
            return cast(
                bytes,
                pandoc.write(
                    doc,
                    format=output_format,
                    options=["--pdf-engine=xelatex"],
                ),
            )
        return cast(bytes, pandoc.write(doc, format=output_format))
    except Exception as exc:
        if output_format == "pdf":
            raise RuntimeError(
                "Falha ao gerar PDF. "
                + "Certifique-se de que o xelatex está instalado "
                + "(apt install texlive-xetex) ou escolha outro formato."
            ) from exc
        raise RuntimeError(
            f"pandoc write failed for format '{output_format}': {exc}"
        ) from exc


def generate_document(
    content: str, title: str, output_format: str
) -> tuple[bytes, str]:
    """Generate a document in the specified format.

    Returns a tuple of (file_bytes, filename_suffix).
    """
    processor = get_processor(output_format)
    return processor.generate(content, title), processor.extension
